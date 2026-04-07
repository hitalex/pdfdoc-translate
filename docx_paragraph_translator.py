# -*- coding: utf-8 -*-
"""
DOCX 段落翻译器
读取已有 docx 文件，将段落和表格单元格内容批量翻译，保留原有格式。
"""

from typing import List, Tuple
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from translator import Translator


def _para_text(para) -> str:
    """拼接段落所有 run 的文本"""
    return "".join(run.text for run in para.runs)


def _set_para_translated(para, new_text: str):
    """
    将翻译文本写回段落。
    保留第一个 run 的格式（字号、加粗、斜体、字体），清空其余 run。
    """
    runs = para.runs
    if not runs:
        run = para.add_run(new_text)
        _apply_cn_font(run)
        return

    # 第一个 run 保留格式，写入翻译文本
    runs[0].text = new_text
    _apply_cn_font(runs[0])

    # 清空其余 run
    for run in runs[1:]:
        run.text = ""


def _apply_cn_font(run, font_name: str = "宋体"):
    """为 run 添加东亚字体，确保中文正确显示"""
    rpr = run._r.get_or_add_rPr()
    # 检查是否已有 rFonts
    existing = rpr.find(qn("w:rFonts"))
    if existing is None:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), font_name)
        rpr.insert(0, rFonts)
    else:
        existing.set(qn("w:eastAsia"), font_name)


def collect_translatable_paragraphs(doc: Document) -> List:
    """
    收集文档中所有需要翻译的段落对象（包括正文段落和表格单元格中的段落）。
    返回段落对象列表（python-docx Paragraph 对象）。
    """
    paras = []

    # 正文段落
    for para in doc.paragraphs:
        if _para_text(para).strip():
            paras.append(para)

    # 表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _para_text(para).strip():
                        paras.append(para)

    return paras


def translate_docx_file(
    input_path: str,
    output_path: str,
    translator: Translator,
    batch_size: int = 20,
) -> None:
    """
    读取 input_path 中的 docx，将所有段落/表格单元格批量翻译后保存到 output_path。

    Args:
        input_path:  原始（OCR重建）docx 路径
        output_path: 翻译后 docx 输出路径
        translator:  已初始化的 Translator 实例
        batch_size:  每批翻译的段落数（过大可能超 token 限制）
    """
    doc = Document(input_path)
    paras = collect_translatable_paragraphs(doc)

    total = len(paras)
    print(f"  共 {total} 个可翻译段落")

    translated_count = 0
    for start in range(0, total, batch_size):
        batch = paras[start: start + batch_size]
        texts = [_para_text(p) for p in batch]

        translated = translator.translate_texts_batch(texts)

        for para, t_text in zip(batch, translated):
            _set_para_translated(para, t_text)

        translated_count += len(batch)
        print(f"  已翻译: {translated_count}/{total} 段")

    doc.save(output_path)
    print(f"  翻译后 DOCX 已保存: {output_path}")
