# -*- coding: utf-8 -*-
"""
创建翻译版DOCX：ISI2025论文"基于主题级偏好优化的多样化论点生成"
保持IEEE双栏格式，图片不翻译，表格内容翻译
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 输出路径 ──────────────────────────────────────────────────────────────────
OUT_PATH = r"D:\code-projects\doc-translate\ISI2025-曹艺琳-中文译版.docx"
IMG_DIR  = r"D:\code-projects\doc-translate\images"


# ══════════════════════════════════════════════════════════════════════════════
#  辅助函数
# ══════════════════════════════════════════════════════════════════════════════

def set_run_font(run, name_en="Times New Roman", name_cn="宋体", size_pt=10,
                 bold=False, italic=False, color=None):
    run.bold  = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = name_en
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name_cn)
    rpr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=3,
             line_spacing=None, first_line_indent=None):
    pf = para.paragraph_format
    pf.alignment      = align
    pf.space_before   = Pt(space_before)
    pf.space_after    = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing      = Pt(line_spacing)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)


def add_paragraph(doc, text, style_name="Normal",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=10, bold=False, italic=False,
                  space_before=0, space_after=3,
                  first_indent=None, color=None,
                  cn_font="宋体", en_font="Times New Roman"):
    para = doc.add_paragraph(style=style_name)
    para_fmt(para, align=align, space_before=space_before,
             space_after=space_after, first_line_indent=first_indent)
    run = para.add_run(text)
    set_run_font(run, name_en=en_font, name_cn=cn_font,
                 size_pt=size, bold=bold, color=color)
    return para


def add_heading(doc, text, level_text="", size=9, bold=True,
                space_before=4, space_after=2,
                align=WD_ALIGN_PARAGRAPH.LEFT):
    """添加节标题（罗马数字格式）"""
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=align,
             space_before=space_before, space_after=space_after)
    run = para.add_run(text)
    set_run_font(run, name_en="Times New Roman", name_cn="黑体",
                 size_pt=size, bold=bold)
    return para


def add_subheading(doc, text, size=9, bold=True,
                   space_before=3, space_after=1):
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=space_before, space_after=space_after)
    run = para.add_run(text)
    set_run_font(run, name_en="Times New Roman", name_cn="黑体",
                 size_pt=size, bold=bold, italic=True)
    return para


def _replace_cols(section, num, space="720", equal_width="1"):
    """移除所有已有的cols元素，添加新的"""
    sectPr = section._sectPr
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for existing in sectPr.findall(f"{{{NS}}}cols"):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), space)
    if num > 1:
        cols.set(qn("w:equalWidth"), equal_width)
    # 插入到sectPr末尾前（sectPr内容顺序灵活）
    sectPr.append(cols)


def set_two_columns(section):
    """将节设为双栏"""
    _replace_cols(section, 2)


def set_single_column(section):
    """将节设为单栏"""
    _replace_cols(section, 1)


def set_page_margins(section,
                     top=1.0, bottom=1.0, left=0.75, right=0.75):
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


def add_column_break(doc):
    """在当前段落末尾插入分栏符"""
    para = doc.add_paragraph(style="Normal")
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "column")
    run._r.append(br)
    return para


def insert_image(doc, img_path, width_inches=3.0, caption=None,
                 caption_size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    """插入图片并可选添加图注"""
    if not os.path.exists(img_path):
        add_paragraph(doc, f"[图片未找到: {img_path}]",
                      align=align, size=8, color=(200, 0, 0))
        return
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=align, space_before=2, space_after=2)
    run = para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(style="Normal")
        para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=1, space_after=4)
        r = cap.add_run(caption)
        set_run_font(r, size_pt=caption_size,
                     name_cn="宋体", name_en="Times New Roman")


def make_table(doc, headers, rows, caption=None,
               font_size=8, bold_header=True):
    """创建带表头的表格"""
    if caption:
        cap = doc.add_paragraph(style="Normal")
        para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=4, space_after=2)
        r = cap.add_run(caption)
        set_run_font(r, size_pt=font_size, bold=True,
                     name_cn="黑体", name_en="Times New Roman")

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=1, space_after=1)
        r = p.add_run(hdr)
        set_run_font(r, size_pt=font_size, bold=bold_header,
                     name_cn="黑体", name_en="Times New Roman")

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=1, space_after=1)
            r = p.add_run(str(val))
            set_run_font(r, size_pt=font_size,
                         name_cn="宋体", name_en="Times New Roman")

    # 设置列宽
    total_width = Inches(6.5)
    col_width = total_width / len(headers)
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_width

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
#  主体：构建文档
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── 第一节：单栏（页眉 + 标题 + 作者 + 摘要）────────────────────────────────
sec0 = doc.sections[0]
set_page_margins(sec0)
set_single_column(sec0)

# 会议页眉
conf_para = doc.add_paragraph(style="Normal")
para_fmt(conf_para, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=0, space_after=2)
r = conf_para.add_run(
    "2025 IEEE国际智能与安全信息学会议（ISI）\n"
    "2025年7月12-14日，中国香港"
)
set_run_font(r, size_pt=8, name_cn="宋体")

# 分隔线
doc.add_paragraph(style="Normal").paragraph_format.space_after = Pt(2)

# 论文标题
title_para = doc.add_paragraph(style="Normal")
para_fmt(title_para, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=6, space_after=4)
r = title_para.add_run("基于主题级偏好优化的多样化论点生成")
set_run_font(r, size_pt=18, bold=True,
             name_cn="黑体", name_en="Times New Roman")

# 作者
authors_para = doc.add_paragraph(style="Normal")
para_fmt(authors_para, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=2, space_after=1)
r = authors_para.add_run(
    "曹艺琳¹·²，张睿珂¹·²，孔庆超¹·²，邢海文³，毛文骥¹·²·*"
)
set_run_font(r, size_pt=10, name_cn="宋体")

# 机构
aff_para = doc.add_paragraph(style="Normal")
para_fmt(aff_para, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=1, space_after=1)
r = aff_para.add_run(
    "¹中国科学院自动化研究所，多模态人工智能系统全国重点实验室\n"
    "²中国科学院大学人工智能学院\n"
    "³国网山东省电力公司"
)
set_run_font(r, size_pt=9, name_cn="宋体")

email_para = doc.add_paragraph(style="Normal")
para_fmt(email_para, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=1, space_after=4)
r = email_para.add_run(
    "{caoyilin2022, zhangruike2020, qingchao.kong, wenji.mao}@ia.ac.cn    "
    "xinghaiwen@sd.sgcc.com.cn"
)
set_run_font(r, size_pt=8, name_cn="宋体", name_en="Courier New")

# 摘要
abs_label = doc.add_paragraph(style="Normal")
para_fmt(abs_label, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=4, space_after=1)
r = abs_label.add_run("摘要")
set_run_font(r, size_pt=9, bold=True, name_cn="黑体")

abs_text = doc.add_paragraph(style="Normal")
para_fmt(abs_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=0, space_after=3, first_line_indent=14)
r = abs_text.add_run(
    "论点生成是社交媒体计算中的核心研究任务，可促进多种安全相关应用。"
    "其旨在自动生成具有说服力且连贯的论点，以表达对给定主张的支持或反对立场。"
    "以往研究主要关注方面控制和事实准确性等问题，在很大程度上忽视了论点生成中的多样性问题。"
    "生成多样化的论点对于帮助用户表达广泛的个人观点至关重要。"
    "随着大语言模型（LLMs）的快速发展，计算论证已成为促进积极在线互动、"
    "推动参与者之间开展积极社区交流的重要手段。然而，这同时也带来了产生有害内容的风险。"
    "因此，在伦理考量下生成多样化论点是另一个重要议题。"
    "本文提出了一种多样化论点生成（DAG）框架来解决上述问题。"
    "DAG将安全感知的监督微调与多样性驱动的数据构建和主题级偏好优化相结合，"
    "鼓励生成主张立场相关且符合伦理的论点，同时兼顾主题多样性和语义多样性。"
    "具体而言，DAG首先对模型进行微调用于论点生成，并使用微调后的模型在符合伦理约束的前提下生成论点；"
    "然后基于主题建模，以主题相关性和语义多样性为指导构建成对偏好数据；"
    "最后通过主题级偏好优化来增强多样化论点生成。"
    "实验结果表明，DAG在生成多样化和可控论点方面取得了比基线更优的性能，同时保持了相当的文本质量。"
)
set_run_font(r, size_pt=9, name_cn="宋体")

idx_para = doc.add_paragraph(style="Normal")
para_fmt(idx_para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=2, space_after=6)
r1 = idx_para.add_run("索引词——")
set_run_font(r1, size_pt=9, bold=True, name_cn="黑体")
r2 = idx_para.add_run("多样化文本生成、论点生成、主题级偏好优化。")
set_run_font(r2, size_pt=9, name_cn="宋体")

# ── 第二节：双栏正文 ──────────────────────────────────────────────────────────
new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
set_page_margins(new_sec)
set_two_columns(new_sec)

BODY_SIZE  = 9
BODY_FONT  = "宋体"
INDENT     = 14     # 首行缩进（磅）

def body(text, indent=INDENT, space_after=3, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=align,
             space_before=0, space_after=space_after,
             first_line_indent=indent if indent else None)
    r = para.add_run(text)
    set_run_font(r, size_pt=BODY_SIZE, name_cn=BODY_FONT)
    return para

def formula(text, space_after=3):
    """公式段落（居中，等宽字体近似）"""
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=2, space_after=space_after)
    r = para.add_run(text)
    set_run_font(r, size_pt=BODY_SIZE, name_cn=BODY_FONT,
                 name_en="Courier New", italic=True)
    return para

def bullet(text, indent_left=18):
    """项目符号段落"""
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=2)
    pf = para.paragraph_format
    pf.left_indent       = Pt(indent_left)
    pf.first_line_indent = Pt(-10)
    r = para.add_run("• " + text)
    set_run_font(r, size_pt=BODY_SIZE, name_cn=BODY_FONT)
    return para


# ────────────────────────────────────────────────────────────
# 第一章  引言
# ────────────────────────────────────────────────────────────
add_heading(doc, "一、引言",
            size=9, bold=True, space_before=4, space_after=2)

body(
    "论点生成是社交媒体计算中的核心研究任务，可促进多种安全相关应用。"
    "其旨在自动生成具有说服力且连贯的论点，以表达对给定主张的支持或反对立场[1]–[3]。"
    "为确保论点生成的质量，生成的论点应与立场保持一致，并与给定主张高度相关。"
    "大多数相关研究集中于论点生成的任务形式化，以满足实际应用需求，包括二元立场控制[1]–[3]、"
    "方面控制[4]、用户个性化[5][6]和事实准确性[7]。然而，以往研究在很大程度上忽视了"
    "论点生成中多样性这一重要问题。"
)
body(
    "在多视角下生成多样化论点是在线互动中的一个重要议题，有助于用户表达广泛的个人观点，"
    "反映社交媒体讨论的丰富性。例如，在图1中，存在四个论点表达对主张"
    "\u300c生成式AI应该免费向公众开放\u300d的反对立场，"
    "涵盖信息操控、算法偏见、隐私风险和网络安全威胁等视角。"
    "相比之下，低多样性论点仅围绕一个视角：虚假信息的生成。"
    "这一示例表明，论点的多样性能有效提升社交媒体讨论质量，避免观点趋同。"
)
body(
    "随着大语言模型（LLMs）的快速发展，计算论证已成为各领域的重要工具，"
    "为增强对有争议主张的理解提供了与特定立场相符的说服性论点。"
    "然而，这同时也带来了产生有害内容的风险。例如，在在线社交媒体平台上，"
    "攻击性、偏颇性、恶意或有害的讨论可能扰乱公共话语，并以负面方式影响公众舆论。"
    "因此，在伦理考量下生成多样化论点是另一个重要议题。"
)
body(
    "为在伦理考量下生成多样化论点，本文提出了多样化论点生成（DAG）框架，"
    "该框架集成了安全感知的监督微调（SFT）、多样性驱动的偏好数据构建和主题级偏好优化。"
    "具体而言，DAG框架首先通过安全感知的监督微调对LLM进行适应，"
    "使基于LLM的论点生成在伦理约束下既保持立场一致性，又与主张高度相关。"
    "然后，DAG通过主题建模，以主题相关性和语义多样性为指导构建成对偏好数据。"
    "最后，微调后的模型通过主题级偏好优化进一步优化，从而在伦理对齐的同时"
    "共同增强主题多样性和语义多样性。"
    "实验结果表明，与基线相比，DAG在生成多样化和可控论点方面取得了显著更优的性能，"
    "同时保持了相当的文本质量。"
)
body("综上，本文的贡献如下：", indent=0)
bullet(
    "在伦理考量下引入多样化论点生成任务，并提出新颖框架DAG来应对其独特的研究挑战。"
)
bullet(
    "DAG将安全感知的SFT与多样性驱动的数据构建和主题级偏好优化相结合，"
    "实现了在主题多样性和语义多样性下的主张立场相关且符合伦理的论点生成。"
)
bullet(
    "实验研究表明，所提出的框架在生成多样化和可控论点方面持续优于竞争性模型，"
    "同时保持了相当的文本质量。"
)

# 图1（位于第一栏）
insert_image(
    doc,
    os.path.join(IMG_DIR, "figure1.png"),
    width_inches=3.0,
    caption=(
        "图1. 具有不同视角的多样化论点示例。在该示例中，"
        "对于主张「生成式AI应免费向公众开放」及其反对立场，"
        "高多样性论点反映了四种不同视角：信息操控、算法偏见、"
        "隐私风险和网络安全威胁。相比之下，低多样性论点仅围绕"
        "一个视角：虚假信息的生成。"
    )
)

# ────────────────────────────────────────────────────────────
# 第二章  相关工作
# ────────────────────────────────────────────────────────────
add_heading(doc, "二、相关工作",
            size=9, bold=True, space_before=4, space_after=2)

add_subheading(doc, "A. 论点生成", size=9)
body(
    "论点生成旨在自动生成具有说服力且连贯的论点，以支持或反对某一主张[8]。"
    "主要方法可大致分为两类：基于检索的方法和基于生成的方法。"
    "大多数相关研究在这两种基础方法之上进一步修改了任务形式化。"
    "[1]首先介绍了一种基于句子检索的辩论论点生成系统。"
    "[9]引入了一种新任务，使用seq2seq框架为给定主张生成多样化的句子论点，"
    "并为每个视角引入潜在机制；然而该方法主要关注论点多样性，忽略了立场控制。"
    "[2]设计了一种基于GPT-2的论点生成流程。"
    "[4]将受控文本生成的概念应用于论点生成领域，Arg-CTRL方法实现了"
    "以给定主张、立场和方面为条件的细粒度控制。"
    "[3]研究了与论证相关的知识图谱在控制论点生成中的应用，并在GPT-2上进行微调。"
    "少数研究关注用户个性化建模[5][6]。为生成高质量的事实性论点，"
    "[7]提出了一个遵循特定立场和论证方案的论点生成器。"
    "[10]提出了论证性文章生成（AEG），旨在生成关于有争议话题或辩论的文章。"
)
body(
    "现有方法主要从特定方面增强论点生成，如立场控制、事实准确性或通过知识图谱、"
    "方面或方案作为额外输入来增强多样性。然而，它们忽视了对主张的语义多样性的考量，"
    "而语义多样性有助于对有争议主张进行更深入的理解。"
)

add_subheading(doc, "B. 文本生成中的多样性", size=9)
body(
    "各种方法被开发用于提升不同文本生成任务中的文本多样性。"
    "混合专家模型被用于增强多样性：[11]进行了实证研究，表明某些类型的混合模型"
    "在翻译任务中提供了文本质量与多样性之间的权衡；"
    "[12]提出了等规模硬期望最大化算法，并训练了多解码器模型用于多样化对话生成。"
    "基于采样的解码方法也被证明能有效提升文本多样性：[13]指出基于最大化的解码方法"
    "会导致重复循环，并提出核采样；[14]通过在MBR解码中引入多样性目标来"
    "开发多样性促进解码算法。TILGAN[15]在潜空间中结合了Transformer自编码器和GAN。"
    "[16]使用上下文学习来多样化LLMs在常识生成任务中的生成结果。"
)
body(
    "现有研究在从不同维度增强多样性方面取得了显著进展。然而，以往研究在很大程度上"
    "忽视了论点生成中的语义多样性问题，而这对于产生细致入微且信息丰富的内容至关重要。"
)

# ────────────────────────────────────────────────────────────
# 第三章  方法
# ────────────────────────────────────────────────────────────
add_heading(doc, "三、方法",
            size=9, bold=True, space_before=4, space_after=2)

body(
    "图2展示了所提出的多样化论点生成框架DAG的整体结构，由安全感知的监督微调、"
    "多样性驱动的偏好数据构建和主题级偏好优化三部分组成。"
)
body(
    "首先，框架采用带有伦理指令的安全感知监督微调，在安全约束下生成与给定主张相关"
    "且与指定立场一致的论点。微调后的模型被用于为每个主张-立场对生成带有相同伦理指令的论点。"
)
body(
    "其次，框架通过主题建模和偏好数据选择进行多样性驱动的偏好数据构建。"
    "对于每个主张-立场对，主题建模从与该主张-立场对相关的生成论点中挖掘代表性主题。"
    "然后，在每个主题内，偏好数据选择基于主题相关性和语义多样性对生成论点进行排序。"
)
body(
    "第三，应用主题级偏好优化进一步优化微调模型以实现多样化论点生成。"
    "它利用构建的偏好数据，增强模型输出与多样性特性（含主题多样性和语义多样性）的对齐。"
)

# 图2（跨栏）
insert_image(
    doc,
    os.path.join(IMG_DIR, "figure2.png"),
    width_inches=6.3,
    caption=(
        "图2. 所提出的多样化论点生成框架DAG的整体结构。"
    )
)

add_subheading(doc, "A. 安全感知的监督微调", size=9)
body(
    "传统的文本生成监督微调使用标注数据集D对LLMs进行对齐，可以表示为："
)
formula("L_SFT(θ) = −E_(x,y)~D [log π_θ(y|x)]          (1)")
body(
    "其中x表示文本生成的指令，y表示生成的文本，θ指模型参数。"
)
body(
    "对于多样化论点生成，给定有争议的主张c∈C和立场s∈S，目标是生成一组"
    "既与立场一致又与主张相关的论点段落Y={y₁,...,yₙ}。"
    "论点段落yᵢ由若干不同论点组成，每个论点通常包含一到两个句子。"
    "为在微调中明确引入立场控制，我们将标准SFT目标扩展为："
)
formula(
    "L_SFT^stance(θ) = −E_(x,y,c,s)~D [log π_θ(y|c,s,x)]   (2)"
)
body(
    "其中LLM被训练为以主张c和立场s为条件生成论点，确保生成输出的立场一致性。"
)
body(
    "安全感知的监督微调阶段结束后，使用得到的立场感知LLM为每个（主张，立场）对"
    "生成n个候选论点段落Y={y₁,y₂,...,yₙ}。每个论点段落yᵢ包含多个句子，"
    "表达给定立场对主张的观点。我们设计了鼓励在伦理推理基础上生成多样化输出的指令，"
    "明确不鼓励有毒、偏颇或歧视性语言，以促进安全和包容性话语。"
)

add_subheading(doc, "B. 多样性驱动的偏好数据构建", size=9)
body(
    "多样性驱动的偏好数据构建模块识别生成的候选论点段落中的关键主题级锚点，"
    "并利用它们指导主题多样性和语义多样性偏好对的选择。"
)
body("a) 主题建模：我们首先使用BERTopic[17]对生成的论点段落Y进行主题建模，"
     "将论点聚类并提取代表性话题关键词。对于每个（主张，立场）对，"
     "模型识别前k个主题，每个主题由关键词组T={t₁,t₂,t₃}表示。"
     "这些主题锚点作为论点集合中不同语义方面的代理。", indent=0)
body("b) 偏好数据选择：我们设计了两个用于偏好数据选择的指标：主题相关性和语义多样性。"
     "对于每个主题锚点tⱼ，使用句子嵌入间的余弦相似度计算每个论点yᵢ与tⱼ的相关性：",
     indent=0)
formula("Relevance(yᵢ, tⱼ) = cos(BERT(yᵢ), BERT(tⱼ))   (3)")
body(
    "其中BERT表示句子嵌入模型Sentence-BERT[18]。"
    "同时，使用Sup-SimCSE[19]通过测量论点段落yᵢ与集合中其他候选的最大语义相似度"
    "来评估其内在多样性："
)
formula(
    "Diversity(yᵢ) = 1 − max_(j≠i) cos(SimCSE(yᵢ), SimCSE(yⱼ))   (4)"
)
body(
    "这确保了与其他论点语义重叠较少的论点获得更高的多样性分数。"
)
body("c) 偏好对选择：对于每个主题锚点tⱼ，通过平衡主题相关性和语义多样性计算"
     "每个论点段落yᵢ的综合分数：", indent=0)
formula(
    "Score(yᵢ, tⱼ) = α · Relevance(yᵢ, tⱼ) + (1−α) · Diversity(yᵢ)   (5)"
)
body(
    "其中α∈[0,1]是平衡系数。基于这些分数，我们选择得分最高的论点作为正样本，"
    "得分最低的作为主题tⱼ下的负样本，形成偏好对(y⁺,y⁻)ⱼ。"
    "对前k个主题重复此过程，为每个输入实例产生k个偏好对；在实现中k=3，"
    "每个输入共产生3个偏好对。正负样本可能在主题间重叠，反映了多样性的多方面性质。"
)
body(
    "通过在数据集中所有（主张，立场）对上应用此策略，我们收集了大量基于主题级语义的偏好对。"
    "DAG通过从候选输出中提取主题锚点，实现结构化的语义比较并指导多样化偏好对的选择，"
    "减少了冗余，最小化了人工工作量，更好地反映了人类偏好，"
    "最终为生成多样化且信息丰富的论点提供了更强的学习信号。"
)

add_subheading(doc, "C. 主题级偏好优化", size=9)
body(
    "为提升论点生成中的语义多样性，我们将直接偏好优化（DPO）[20]应用于"
    "构建的多样性驱动偏好数据集。给定主张c和立场s，偏好数据集定义为："
)
formula(
    "D_pref = {(c, s, y⁺, y⁻) | φ(y⁺) > φ(y⁻)}   (6)"
)
body(
    "其中y⁺和y⁻分别表示在相同（c,s）对下语义多样性更高和更低的论点，"
    "φ(·)是多样性评分函数。DPO目标通过比较策略模型πθ与参考模型πref（通常为SFT模型）"
    "的输出概率来使模型与人类偏好信号对齐。目标函数定义为："
)
formula(
    "L(πθ; πref) = −E[log σ(β log(πθ(y⁺)/πref(y⁺)) − β log(πθ(y⁻)/πref(y⁻)))]   (7)"
)
body("其中：", indent=0)
bullet("πθ(y|c,s,x)是在立场约束下优化语义多样性的策略模型（LLM）；")
bullet("πref(y|c,s,x)是参考模型，通常是初始SFT模型；")
bullet("β是控制偏好学习强度的温度缩放因子；")
bullet("σ(·)是Sigmoid函数，确保优化过程的稳定性。")
body(
    "偏好优化鼓励模型为具有更强语义多样性的论点分配更高的概率，"
    "同时对重复或趋同的输出进行惩罚。"
)

# ────────────────────────────────────────────────────────────
# 第四章  实验
# ────────────────────────────────────────────────────────────
add_heading(doc, "四、实验",
            size=9, bold=True, space_before=4, space_after=2)

add_subheading(doc, "A. 数据集", size=9)
body(
    "我们使用PERSPECTRUM数据集[21]训练和评估多样化论点生成模型。"
    "原始数据集包含907个主张，每个主张与从在线辩论论坛和网络中收集的"
    "多个带立场标注的论点相关联。每条数据由（主张、立场、视角、证据）四元组组成，"
    "其中视角是表达支持或反对主张观点的简短句子。"
    "对于每个（主张，立场）对，我们将所有相关视角拼接为段落作为生成目标。"
)
body(
    "为评估模型在未见过的多样主张上的泛化能力，我们进一步构建了一个混合数据集用于评估，"
    "纳入来自多个领域外立场检测数据集的主张。具体而言，我们从四个广泛使用的基准数据集中"
    "采样主张：WT-WT[22]、P-Stance[23]、SemEval-2016[24]和COVID19-Stance[25]。"
    "这些数据集涵盖政治、社会议题和公共卫生等广泛领域，"
    "为评估超越PERSPECTRUM领域的论点生成提供了综合测试平台。"
)

add_subheading(doc, "B. 评估指标", size=9)
body(
    "我们从四个角度评估生成的论点：多样性、可控性、文本质量和安全性。"
    "对于词汇多样性和流畅性的自动评估，我们计算Dist-n分数[26]"
    "（Dist-1、Dist-2、Dist-3和平均Dist），用于衡量生成输出中唯一n-gram的比例。"
    "我们还报告使用GPT-2 medium[27]计算的困惑度（PPL），"
    "作为流畅性和语法质量的指标。"
)
body(
    "为评估可控性，我们使用GPT-4o-mini作为评估器，"
    "对每个输出在主张相关性和立场对齐性方面进行0–5分制评分。"
    "相关性衡量论点与主张的对应程度；立场对齐性反映与指定立场的一致程度。"
    "我们进行人工评估以评估生成论点的安全性。"
)

add_subheading(doc, "C. 基线方法", size=9)
body(
    "我们在Meta-Llama3-8B-Instruct和Qwen2.5-7B-Instruct上评估所提出的DAG框架。"
    "我们将结果与涵盖不同学习设置的几种基线进行比较："
    "零样本提示、少样本提示、监督微调（SFT）、GPT-3.5-Turbo和GPT-4-1106-Preview提示。"
    "在基于提示的设置中，模型通过有无上下文示例的手动指令进行引导。"
    "在微调设置中，模型在我们的立场标注数据集上训练，以提升可控性和多样性。"
    "我们还包括一个指令微调基线，作为策略级消融，"
    "以评估指令微调、微调和所提出方法在多样性、可控性和文本质量方面的各自影响。"
)

add_subheading(doc, "D. 主要结果与分析", size=9)
body(
    "表I和表II报告了在领域内数据集和领域外数据集上的实验结果。"
)

# 表I
make_table(
    doc,
    headers=[
        "模型", "方法",
        "立场对齐↑", "主张相关↑",
        "Dist-1↑", "Dist-2↑", "Dist-3↑", "均Dist↑",
        "PPL↓"
    ],
    rows=[
        ["LLaMA3-8B", "零样本提示",     "4.82","4.35","0.10","0.45","0.72","0.42","9.82"],
        ["",          "少样本提示",     "4.71","4.66","0.13","0.39","0.56","0.36","55.47"],
        ["",          "监督微调",       "4.94","4.41","0.11","0.45","0.72","0.43","9.59"],
        ["",          "GPT-3.5-Turbo", "4.99","4.02","0.10","0.46","0.74","0.44","11.79"],
        ["",          "GPT-4-1106",    "4.93","4.56","0.11","0.46","0.73","0.43","15.49"],
        ["",          "DAG(LLaMA3-8B)","4.94","4.97","0.12","0.48","0.74","0.45","9.81"],
        ["Qwen2.5-7B","零样本提示",     "4.90","4.29","0.12","0.52","0.80","0.48","13.45"],
        ["",          "少样本提示",     "4.97","4.80","0.11","0.46","0.70","0.43","20.62"],
        ["",          "监督微调",       "4.91","4.37","0.12","0.52","0.79","0.48","12.98"],
        ["",          "GPT-3.5-Turbo", "4.99","4.02","0.10","0.46","0.74","0.44","11.79"],
        ["",          "GPT-4-1106",    "4.93","4.56","0.11","0.46","0.73","0.43","15.49"],
        ["",          "DAG(Qwen2.5-7B)","4.95","4.96","0.13","0.56","0.83","0.51","14.02"],
    ],
    caption="表I  PERSPECTRUM数据集上多样化论点生成实验结果",
    font_size=7,
)

body("", space_after=2)

# 表II
make_table(
    doc,
    headers=[
        "模型", "方法",
        "立场对齐↑", "主张相关↑",
        "Dist-1↑", "Dist-2↑", "Dist-3↑", "均Dist↑",
        "PPL↓"
    ],
    rows=[
        ["LLaMA3-8B", "零样本提示",      "4.78","4.06","0.20","0.52","0.68","0.47","14.85"],
        ["",          "少样本提示",      "5.00","5.00","0.06","0.13","0.16","0.12","29.02"],
        ["",          "监督微调",        "4.92","3.81","0.21","0.54","0.69","0.48","14.89"],
        ["",          "GPT-3.5-Turbo",  "4.86","4.58","0.26","0.68","0.86","0.60","11.36"],
        ["",          "GPT-4-1106",     "4.90","4.72","0.28","0.68","0.82","0.59","25.96"],
        ["",          "DAG(LLaMA3-8B)", "5.00","4.95","0.26","0.60","0.73","0.53","15.25"],
        ["Qwen2.5-7B","零样本提示",      "4.33","3.08","0.27","0.71","0.89","0.62","13.20"],
        ["",          "少样本提示",      "5.00","5.00","0.29","0.68","0.83","0.60","19.40"],
        ["",          "监督微调",        "4.33","3.53","0.27","0.72","0.89","0.62","12.89"],
        ["",          "GPT-3.5-Turbo",  "4.86","4.58","0.26","0.68","0.86","0.60","11.36"],
        ["",          "GPT-4-1106",     "4.90","4.72","0.28","0.68","0.82","0.59","25.96"],
        ["",          "DAG(Qwen2.5-7B)","4.97","4.80","0.32","0.76","0.91","0.66","13.07"],
    ],
    caption="表II  混合数据集上多样化论点生成实验结果",
    font_size=7,
)

body("", space_after=2)

body(
    "在领域内PERSPECTRUM数据集（如表I所示）上，DAG在可控性和文本多样性方面"
    "取得了最佳综合性能，同时保持了相当的文本质量，"
    "证明了其在伦理考量下生成连贯且多样化论点的能力。"
)
body(
    "在领域外混合数据集（如表II所示）上，少样本提示方法产生了较强的可控性分数，"
    "但往往以牺牲文本多样性（如LLaMA3-8B）或质量（如LLaMA3-8B和Qwen2.5-7B）为代价。"
    "相比之下，DAG持续取得更平衡的性能，在不损害文本质量的前提下增强了多样性和可控性。"
)
body(
    "总体而言，DAG在生成多样化论点方面表现出优越性能。"
    "与基线方法相比，DAG持续提升多样性，同时在可控性（立场对齐和主张相关性）"
    "和文本质量（PPL）方面保持强劲性能。DAG突显了主题级偏好优化在增强论点生成"
    "内容丰富性和可控性方面的有效性。此外，DAG通过利用可解释的主题级锚点提升了透明度，"
    "有助于社会计算分析并减少潜在误用，降低了对黑盒LLM行为的过度依赖。"
)

add_subheading(doc, "E. 模型安全性的人工评估", size=9)
body(
    "为评估模型的安全性，我们基于混合数据集使用两名独立评判员进行人工评估。"
    "他们被指导从伦理角度对论点进行标注，识别恶意、有毒、偏颇或歧视性的输出。"
    "评判员1和评判员2分别取得0.958和0.972的准确率分数，"
    "平均比率为0.965，评判员间一致性为0.793（Cohen's Kappa）。"
)

# ────────────────────────────────────────────────────────────
# 第五章  结论
# ────────────────────────────────────────────────────────────
add_heading(doc, "五、结论",
            size=9, bold=True, space_before=4, space_after=2)
body(
    "本文提出了DAG，一种在伦理考量下进行多样化论点生成的框架。"
    "该框架集成了安全感知的监督微调、多样性驱动的偏好数据构建和主题级偏好优化，"
    "以支持在伦理约束下生成主张立场相关、主题感知、多样性增强的论点。"
    "实验结果表明，与竞争性基线模型相比，该框架在生成多样化和可控论点方面"
    "取得了相当更优的性能，同时保持了相当的文本质量。"
    "本工作强调了多样性和主题建模在论点生成中的重要性，"
    "并提供了一个基于LLM的框架，用于生成语义和主题上多样化的论点，支持更安全的在线话语。"
)

# ────────────────────────────────────────────────────────────
# 致谢
# ────────────────────────────────────────────────────────────
add_heading(doc, "致谢",
            size=9, bold=True, space_before=4, space_after=2)
body(
    "本工作得到国家电网有限公司科技项目资助：电力行业语义大模型构建技术研究"
    "及其在运维智能辅助决策中的应用（5700-202416227A-1-1-ZN）。"
)

# ────────────────────────────────────────────────────────────
# 参考文献
# ────────────────────────────────────────────────────────────
add_heading(doc, "参考文献",
            size=9, bold=True, space_before=4, space_after=2)

refs = [
    "[1] M. Sato et al., “End-to-end argument generation system in debating,” ACL System Demonstrations, 2015.",
    "[2] S. Gretz et al., “The work week is the best time to start a family – a study of gpt-2 based claim generation,” EMNLP Findings, 2020.",
    "[3] K. AlKhatib et al., “Employing argumentation knowledge graphs for neural argument generation,” ACL, 2021.",
    "[4] B. Schiller et al., “Aspect-controlled neural argument generation,” NAACL, 2021.",
    "[5] M. Alshomary et al., “Belief-based generation of argumentative claims,” EACL, 2021.",
    "[6] M. Alshomary et al., “The moral debater: A study on the computational generation of morally framed arguments,” ACL, 2022.",
    "[7] S. Saha and R. Srihari, “ArgU: A controllable factual argument generator,” ACL, 2023.",
    "[8] G. Chen et al., “Exploring the potential of large language models in computational argumentation,” ACL, 2024.",
    "[9] C. Park et al., “Generating sentential arguments from diverse perspectives on controversial topic,” NLP4IF Workshop, 2019.",
    "[10] R. Xiao et al., “Prove your point!: Bringing proof-enhancement principles to argumentative essay generation,” EMNLP, 2024.",
    "[11] T. Shen et al., “Mixture models for diverse machine translation: Tricks of the trade,” ICML, 2019.",
    "[12] Y. Wen et al., “An equal-size hard EM algorithm for diverse dialogue generation,” ICLR, 2023.",
    "[13] A. Holtzman et al., “The curious case of neural text degeneration,” ICLR, 2020.",
    "[14] Y. Jinnai et al., “Generating diverse and high-quality texts by minimum Bayes risk decoding,” ACL Findings, 2024.",
    "[15] S. Diao et al., “TILGAN: Transformer-based implicit latent GAN for diverse and coherent text generation,” ACL Findings, 2021.",
    "[16] T. Zhang et al., “Improving diversity of commonsense generation by large language models via in-context learning,” EMNLP Findings, 2024.",
    "[17] M. Grootendorst, “BERTopic: Neural topic modeling with a class-based tf-idf procedure,” arXiv:2203.05794, 2022.",
    "[18] N. Reimers and I. Gurevych, “Sentence-BERT: Sentence embeddings using Siamese BERT-networks,” EMNLP-IJCNLP, 2019.",
    "[19] T. Gao et al., “SimCSE: Simple contrastive learning of sentence embeddings,” EMNLP, 2021.",
    "[20] R. Rafailov et al., “Direct preference optimization: Your language model is secretly a reward model,” NeurIPS, vol. 36, 2024.",
    "[21] S. Chen et al., “Seeing things from a different angle: Discovering diverse perspectives about claims,” NAACL, 2019.",
    "[22] C. Conforti et al., “Will-they-won't-they: A very large dataset for stance detection on Twitter,” ACL, 2020.",
    "[23] Y. Li et al., “P-stance: A large dataset for stance detection in political domain,” ACL Findings, 2021.",
    "[24] S. Mohammad et al., “SemEval-2016 task 6: Detecting stance in tweets,” SemEval, 2016.",
    "[25] K. Glandt et al., “Stance detection in COVID-19 tweets,” ACL-IJCNLP, 2021.",
    "[26] J. Li et al., arXiv:1510.03055, 2015.",
    "[27] A. Radford et al., “Language models are unsupervised multitask learners,” OpenAI blog, 2019.",
]

for ref in refs:
    para = doc.add_paragraph(style="Normal")
    para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=1)
    pf = para.paragraph_format
    pf.left_indent       = Pt(12)
    pf.first_line_indent = Pt(-12)
    r = para.add_run(ref)
    set_run_font(r, size_pt=8, name_cn="宋体",
                 name_en="Times New Roman")

# ── 保存 ─────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print("Saved: " + OUT_PATH)
