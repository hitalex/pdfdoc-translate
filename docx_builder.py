# -*- coding: utf-8 -*-
"""
DOCX 输出构建器
策略：使用两列 Table 作为布局容器，保留双栏、表格、图片
"""

import io
import os
from typing import List, Optional
from bs4 import BeautifulSoup
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from layout_analyzer import PageLayout, Region, TYPE_TITLE, TYPE_TEXT, \
    TYPE_TABLE, TYPE_FIGURE, TYPE_FIG_CAP, TYPE_TBL_CAP


# ── 尺寸常量（A4）───────────────────────────────────────────────────────────
PAGE_W   = Cm(21.0)
PAGE_H   = Cm(29.7)
MARGIN   = Cm(2.0)
COL_GAP  = Cm(0.5)
CONTENT_W = PAGE_W - MARGIN * 2           # 可用宽度
HALF_COL  = (CONTENT_W - COL_GAP) / 2    # 单栏宽度


class DocxBuilder:
    """将翻译后的版面数据构建为 DOCX 文档"""

    def __init__(self):
        pass

    def build(self, layouts: List[PageLayout], output_path: str,
              verbose: bool = False):
        self._verbose = verbose
        self._write_log: list = []   # 记录每个区域的写入情况

        doc = Document()
        self._setup_page(doc)

        for i, layout in enumerate(layouts):
            if i > 0:
                self._add_page_break(doc)
            self._build_page(doc, layout)

        doc.save(output_path)
        print(f"  DOCX 已保存: {output_path}")

        if verbose:
            print("\n  ── DOCX 写入日志 ──")
            for entry in self._write_log:
                print(f"  {entry}")
            written = sum(1 for e in self._write_log if not e.startswith("  SKIP"))
            skipped = sum(1 for e in self._write_log if e.startswith("  SKIP"))
            print(f"  共写入 {written} 个区域，跳过 {skipped} 个")

    # ── 页面设置 ─────────────────────────────────────────────────────────────

    def _setup_page(self, doc: Document):
        for section in doc.sections:
            section.page_width  = PAGE_W
            section.page_height = PAGE_H
            section.left_margin   = MARGIN
            section.right_margin  = MARGIN
            section.top_margin    = MARGIN
            section.bottom_margin = MARGIN

    # ── 构建单页 ─────────────────────────────────────────────────────────────

    def _build_page(self, doc: Document, layout: PageLayout):
        regions = layout.get_sorted_regions()

        if not layout.is_two_column:
            # 单栏：直接逐个添加
            for r in regions:
                self._add_region_single(doc, r, CONTENT_W)
        else:
            # 双栏：分组后放入 2 列布局表
            self._build_two_column_page(doc, layout, regions)

    def _build_two_column_page(self, doc: Document, layout: PageLayout,
                                regions: List[Region]):
        """
        双栏页面构建：
        - 全宽区域（full）直接添加到文档
        - 左右栏内容成对放入一个 2 列布局表的行
        """
        full_regions = [r for r in regions if r.column == "full"]
        left_regions  = [r for r in regions if r.column == "left"]
        right_regions = [r for r in regions if r.column == "right"]

        # 全宽内容先输出
        for r in full_regions:
            self._add_region_single(doc, r, CONTENT_W)

        # 若有双栏内容，用 2 列布局表
        if left_regions or right_regions:
            max_rows = max(len(left_regions), len(right_regions))
            tbl = doc.add_table(rows=max_rows, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _set_table_no_border(tbl)
            _set_col_width(tbl, 0, HALF_COL)
            _set_col_width(tbl, 1, HALF_COL)

            for row_idx in range(max_rows):
                row = tbl.rows[row_idx]
                if row_idx < len(left_regions):
                    cell = row.cells[0]
                    self._fill_cell(cell, left_regions[row_idx])
                if row_idx < len(right_regions):
                    cell = row.cells[1]
                    self._fill_cell(cell, right_regions[row_idx])

    def _fill_cell(self, cell, region: Region):
        """向表格单元格中填充区域内容"""
        # 清空默认空段落
        for p in cell.paragraphs:
            p.clear()

        col_w = HALF_COL
        if region.type == TYPE_TABLE:
            self._add_table_to_cell(cell, region)
        elif region.type == TYPE_FIGURE:
            self._add_image_to_para(cell.paragraphs[0], region, col_w)
        else:
            text = region.translated_text or region.text
            p = cell.paragraphs[0]
            self._set_para_text(p, text, is_title=(region.type == TYPE_TITLE))

    # ── 单栏区域添加 ─────────────────────────────────────────────────────────

    def _log(self, region: Region, action: str):
        if hasattr(self, '_verbose') and self._verbose:
            prefix = "  SKIP" if action.startswith("跳过") else "  WRITE"
            self._write_log.append(
                f"{prefix} [{region.type:<14}] {action}"
            )

    def _add_region_single(self, doc: Document, region: Region, col_width):
        if region.type == TYPE_FIGURE:
            if region.image is not None:
                self._log(region, f"插入图片 {region.image.size}")
            else:
                self._log(region, "跳过图片（image=None）")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._add_image_to_para(p, region, col_width)

        elif region.type == TYPE_TABLE:
            text = region.translated_text or region.text
            if text:
                self._add_paragraph(doc, text, is_title=False)
            if region.table_html:
                self._log(region, f"插入表格 HTML({len(region.table_html)}字符)")
                self._add_docx_table(doc, region)
            else:
                self._log(region, "跳过表格（table_html 为空）")

        elif region.type == TYPE_TITLE:
            text = region.translated_text or region.text
            if text:
                self._log(region, f'写入标题: "{text[:60]}"')
                self._add_paragraph(doc, text, is_title=True)
            else:
                self._log(region, "跳过标题（text 为空）")

        else:
            text = region.translated_text or region.text
            if text:
                self._log(region, f'写入文本: "{text[:60]}"')
                self._add_paragraph(doc, text, is_title=False)
            else:
                self._log(region, "跳过文本（text 为空）")

    def _add_paragraph(self, doc: Document, text: str, is_title: bool = False):
        p = doc.add_paragraph()
        self._set_para_text(p, text, is_title=is_title)
        return p

    def _set_para_text(self, p, text: str, is_title: bool = False):
        run = p.add_run(text)
        pf = p.paragraph_format
        if is_title:
            run.bold = True
            run.font.size = Pt(13)
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_after = Pt(6)
        else:
            run.font.size = Pt(10)
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.space_after = Pt(3)
            pf.first_line_indent = Pt(20)
        # 中文字体
        _set_east_asia_font(run, "宋体")

    # ── 图片插入 ─────────────────────────────────────────────────────────────

    def _add_image_to_para(self, p, region: Region, max_width):
        if region.image is None:
            return
        try:
            img_bytes = io.BytesIO()
            region.image.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # 计算显示尺寸（保持宽高比，最大宽度为 max_width）
            orig_w, orig_h = region.image.size
            display_w = min(max_width, Inches(orig_w / 96))  # 96 DPI 基准
            run = p.add_run()
            run.add_picture(img_bytes, width=display_w)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"    图片插入失败: {e}")

    # ── 表格构建 ─────────────────────────────────────────────────────────────

    def _add_docx_table(self, doc: Document, region: Region):
        html = region.translated_html or region.table_html
        if not html:
            return

        soup = BeautifulSoup(html, "html.parser")
        rows_html = soup.find_all("tr")
        if not rows_html:
            return

        # 确定列数
        max_cols = max(
            sum(int(td.get("colspan", 1)) for td in tr.find_all(["td", "th"]))
            for tr in rows_html
        )
        if max_cols == 0:
            return

        tbl = doc.add_table(rows=len(rows_html), cols=max_cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, tr in enumerate(rows_html):
            cells_html = tr.find_all(["td", "th"])
            col_cursor = 0
            for td in cells_html:
                if col_cursor >= max_cols:
                    break
                text = td.get_text(strip=True)
                cell = tbl.rows[row_idx].cells[col_cursor]
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(9)
                _set_east_asia_font(run, "宋体")
                if td.name == "th":
                    run.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                col_cursor += int(td.get("colspan", 1))

        doc.add_paragraph()  # 表格后空行

    def _add_table_to_cell(self, cell, region: Region):
        """在布局表单元格内插入数据表"""
        html = region.translated_html or region.table_html
        if not html:
            return
        soup = BeautifulSoup(html, "html.parser")
        rows_html = soup.find_all("tr")
        if not rows_html:
            return
        max_cols = max(
            sum(int(td.get("colspan", 1)) for td in tr.find_all(["td", "th"]))
            for tr in rows_html
        ) or 1
        tbl = cell.add_table(rows=len(rows_html), cols=max_cols)
        tbl.style = "Table Grid"
        for row_idx, tr in enumerate(rows_html):
            col_cursor = 0
            for td in tr.find_all(["td", "th"]):
                if col_cursor >= max_cols:
                    break
                text = td.get_text(strip=True)
                c = tbl.rows[row_idx].cells[col_cursor]
                run = c.paragraphs[0].add_run(text)
                run.font.size = Pt(8)
                _set_east_asia_font(run, "宋体")
                col_cursor += int(td.get("colspan", 1))

    # ── 分页 ─────────────────────────────────────────────────────────────────

    @staticmethod
    def _add_page_break(doc: Document):
        p = doc.add_paragraph()
        run = p.add_run()
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _el
        br = _el("w:br")
        br.set(_qn("w:type"), "page")
        run._r.append(br)


# ── 辅助函数 ─────────────────────────────────────────────────────────────────

def _set_east_asia_font(run, font_name: str = "宋体"):
    """设置东亚文字字体"""
    rpr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rpr.insert(0, rFonts)


def _set_table_no_border(table):
    """移除表格所有边框（用于布局容器）"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_col_width(table, col_idx: int, width):
    """设置表格列宽"""
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        # width 是 Emu，转为 twips (1 inch = 1440 twips = 914400 Emu)
        twips = int(width / 914400 * 1440)
        tcW.set(qn("w:w"), str(twips))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)
